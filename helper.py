from docx import Document

def merge_cells(doc):
    # Создание таблицы
    table = doc.add_table(rows=5, cols=5)

    # Заполнение ячеек данными
    for row in range(5):
        for col in range(5):
            cell = table.cell(row, col)
            cell.text = f'[{row},{col}]'

    # Объединение ячеек по строкам (ячейка из столбца 0, строки 0-1)
    cell_to_merge_row = table.cell(0, 0)
    cell_to_merge_row.merge(table.cell(1, 0))

    # Объединение ячеек по столбцам (ячейка из первого ряда, столбцы 1-2)
    cell_to_merge_col = table.cell(0, 1)
    cell_to_merge_col.merge(table.cell(0, 2))

    # Объединение ячеек одновременно (ячейка из 0-1 ряда, 3-4 столбца)
    cell_to_merge_both = table.cell(0, 3)
    cell_to_merge_both.merge(table.cell(1, 4))

    # Установка текста для объединенных ячеек
    cell_to_merge_row.text = 'Объединение по строкам'
    cell_to_merge_col.text = 'Объединение по столбцам'
    cell_to_merge_both.text = 'Объединение обеих'

# Создание документа
doc = Document()
doc.add_heading('Пример объединения ячеек', level=1)
merge_cells(doc)

# Сохранение документа
doc.save('merged_cells_example.docx')


